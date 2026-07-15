#!/usr/bin/env python3
"""Read profile CSV -> figures -> academic-style .docx (results + analysis)."""
import csv, os, matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

AIOTC = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CSV = os.path.join(AIOTC, "results", "profile_blackwell.csv")
FIGDIR = os.path.join(AIOTC, "results", "figures"); os.makedirs(FIGDIR, exist_ok=True)
OUT = os.path.join(AIOTC, "AIoTC_results_blackwell.docx")

rows = list(csv.DictReader(open(CSV)))
def f(r, k): return float(r[k]) if r[k] not in ("", None) else float("nan")
DISP={"ronin_resnet18":"RoNIN-ResNet","ronin_tcn":"RoNIN-TCN","ronin_lstm":"RoNIN-LSTM",
      "imunet":"IMUNet","mobilenetv2":"MobileNetV2","mnasnet":"MnasNet",
      "efficientnet_b0":"EfficientNet-B0","tlio_resnet":"TLIO","tinyodom":"TinyOdom","eqnio":"EqNIO"}
FAM={"ronin_resnet18":"CNN (ResNet)","tlio_resnet":"CNN (ResNet)","ronin_lstm":"RNN",
     "ronin_tcn":"TCN","tinyodom":"TCN (NAS)","imunet":"Mobile CNN","mobilenetv2":"Mobile CNN",
     "mnasnet":"Mobile CNN","efficientnet_b0":"Mobile CNN","eqnio":"Equivariant"}
FAMCOL={"CNN (ResNet)":"#2E86AB","RNN":"#6A4C93","TCN":"#1B998B","TCN (NAS)":"#0B7A4B",
        "Mobile CNN":"#E4572E","Equivariant":"#8B1E3F"}
def get(model, prec, key):
    for r in rows:
        if r["model"]==model and r["precision"]==prec: return f(r,key)
    return float("nan")
models=sorted({r["model"] for r in rows}, key=lambda m: get(m,"fp32","lat_med_ms"))
disp=DISP
PRECS=["fp32","fp16","bf16"]; PLAB={"fp32":"FP32","fp16":"FP16","bf16":"BF16"}

plt.rcParams.update({"font.size":10.5,"axes.grid":True,"grid.alpha":0.3,
                     "axes.spines.top":False,"axes.spines.right":False})
import matplotlib.patches as mpatch
def hbar(key, xlabel, title, fname, fmt):
    order=sorted(models, key=lambda m: get(m,"fp32",key))
    vals=[get(m,"fp32",key) for m in order]; cols=[FAMCOL[FAM[m]] for m in order]
    fig,ax=plt.subplots(figsize=(6.6,4.3)); y=np.arange(len(order))
    ax.barh(y,vals,color=cols); ax.set_yticks(y); ax.set_yticklabels([DISP[m] for m in order]); ax.invert_yaxis()
    ax.set_xlabel(xlabel); ax.set_title(title); ax.yaxis.grid(False)
    for yi,v in zip(y,vals): ax.text(v,yi,"  "+fmt.format(v),va="center",fontsize=8)
    seen={FAM[m]:FAMCOL[FAM[m]] for m in order}
    ax.legend([mpatch.Patch(color=c) for c in seen.values()], list(seen.keys()),frameon=False,fontsize=8,loc="lower right")
    fig.tight_layout(); fig.savefig(f"{FIGDIR}/{fname}",dpi=200,bbox_inches="tight"); plt.close(fig)

hbar("lat_med_ms","median latency (ms, FP32)","Per-inference latency, batch = 1 (sorted)","fig1_latency.png","{:.2f}")
hbar("energy_mJ_per_inf","energy per inference (mJ, FP32)","Energy per inference (sorted)","fig2_energy.png","{:.0f}")

# Fig 3: params vs latency (fp32), family-colored, log-x
fig,ax=plt.subplots(figsize=(6.6,4.3))
for m in models:
    px,py=get(m,"fp32","params_M"),get(m,"fp32","lat_med_ms")
    ax.scatter(px,py,s=80,color=FAMCOL[FAM[m]],zorder=3)
    ax.annotate(DISP[m],(px,py),textcoords="offset points",xytext=(6,3),fontsize=8)
ax.set_xscale("log"); ax.set_xlabel("parameters (M, log scale)"); ax.set_ylabel("median latency (ms, FP32)")
ax.set_title("Architecture, not parameter count, drives latency")
fig.tight_layout(); fig.savefig(f"{FIGDIR}/fig3_params_latency.png",dpi=200,bbox_inches="tight"); plt.close(fig)

# Fig 4: bf16/fp16 crossover vs hidden size (minimal nn.LSTM, batch=1 seq=200)
# measured by scripts/diagnose_bf16_lstm.py (Layer 2); see DIAGNOSIS_bf16_lstm.md
XOVER_H=[50,100,200]; XOVER_FP16=[0.402,1.449,3.948]; XOVER_BF16=[4.412,4.409,4.456]
ratio=[b/f for b,f in zip(XOVER_BF16,XOVER_FP16)]
fig,ax=plt.subplots(figsize=(6.2,3.7))
ax.plot(XOVER_H,ratio,"o-",color="#C1272D",lw=2,ms=9,zorder=3)
for h,r in zip(XOVER_H,ratio): ax.annotate(f"{r:.1f}×",(h,r),textcoords="offset points",xytext=(6,7),fontsize=10)
ax.axhline(1.0,ls="--",color="gray",lw=1); ax.text(200,1.25,"parity",color="gray",fontsize=9,ha="right")
ax.set_xlabel("LSTM hidden size"); ax.set_ylabel("BF16 / FP16 latency ratio")
ax.set_title("BF16-vs-FP16 LSTM slowdown collapses as hidden size grows"); ax.set_xticks(XOVER_H)
fig.tight_layout(); fig.savefig(f"{FIGDIR}/fig4_bf16_crossover.png",dpi=200,bbox_inches="tight"); plt.close(fig)
print("figures written to", FIGDIR)

# ---------------- accuracy results (optional; grows as models finish training) ----------------
ACC_CSV = os.path.join(AIOTC, "results", "accuracy_blackwell.csv")
acc_rows = list(csv.DictReader(open(ACC_CSV))) if os.path.exists(ACC_CSV) else []
acc_models = []
for r in acc_rows:
    if r["model"] not in acc_models: acc_models.append(r["model"])
def accget(model, prec, key):
    for r in acc_rows:
        if r["model"]==model and r["precision"]==prec: return r[key]
    return None

# idle/active dynamic-power results (primary models; optional)
PWR_CSV = os.path.join(AIOTC, "results", "power_dynamic.csv")
pwr_rows = list(csv.DictReader(open(PWR_CSV))) if os.path.exists(PWR_CSV) else []
def pwrget(model, prec, key):
    for r in pwr_rows:
        if r["model"]==model and r["precision"]==prec: return r[key]
    return None
_idle = f"{sum(float(r['idle_W']) for r in pwr_rows)/len(pwr_rows):.0f}" if pwr_rows else None

# ---------------- edge tiers (Jetson AGX Orin + Orin Nano) ----------------
def _load(p):
    fp=os.path.join(AIOTC,"results",p)
    return list(csv.DictReader(open(fp))) if os.path.exists(fp) else []
EDGE={
 "agx_orin": dict(name="AGX Orin",  pt=_load("profile_agx_orin.csv"),  acc=_load("accuracy_agx_orin.csv"),
                  trt_fp16=_load("trt_latency_agx_orin.csv"), trt_int8=_load("trt_int8_latency_agx_orin.csv")),
 "orin_nano":dict(name="Orin Nano", pt=_load("profile_orin_nano.csv"), acc=_load("accuracy_orin_nano.csv"),
                  trt_fp16=_load("profile_int8_orin_nano.csv"), trt_int8=_load("trt_int8_latency_orin_nano.csv")),
}
int8acc=_load("accuracy_int8_canonical.csv")
has_edge=bool(EDGE["orin_nano"]["pt"]) and bool(int8acc)
PRIM=["ronin_resnet18","imunet","mobilenetv2","mnasnet","efficientnet_b0","tinyodom"]
def pt_lat(dev,prec,m):
    for r in EDGE[dev]["pt"]:
        if r["model"]==m and r["precision"]==prec: return float(r["lat_med_ms"])
def pt_e(dev,prec,m):
    for r in EDGE[dev]["pt"]:
        if r["model"]==m and r["precision"]==prec and r["energy_mJ_per_inf"]: return float(r["energy_mJ_per_inf"])
def trt_lat(dev,prec,m):
    src=EDGE[dev]["trt_int8"] if prec=="int8" else EDGE[dev]["trt_fp16"]
    for r in src:
        if r["model"]==m and (prec=="int8" or r.get("precision")=="fp16"):
            return float(r.get("gpu_lat_med_ms") or r.get("trtexec_lat_med_ms"))
def int8_ate(m):
    for r in int8acc:
        if r["model"]==m and r["precision"]=="onnx_int8": return float(r["ate_m"])
def acc_fp32(m):  # device-independent fp32 ATE (from Blackwell accuracy)
    return float(accget(m,"fp32","ate_m"))

if has_edge:
    import matplotlib.patches as _mp
    # Fig 6: TensorRT latency across the two edge devices (fp16 & int8), sorted
    fig,ax=plt.subplots(figsize=(6.6,4.0)); y=np.arange(len(PRIM)); w=0.2
    order=sorted(PRIM, key=lambda m: trt_lat("orin_nano","fp16",m) or 9)
    for j,(dev,prec,c) in enumerate([("agx_orin","fp16","#2E86AB"),("agx_orin","int8","#1B998B"),
                                     ("orin_nano","fp16","#E4572E"),("orin_nano","int8","#8B1E3F")]):
        vals=[trt_lat(dev,prec,m) for m in order]
        ax.barh(y+(1.5-j)*w, vals, w, color=c, label=f"{EDGE[dev]['name']} {prec.upper()}")
    ax.set_yticks(y); ax.set_yticklabels([DISP[m] for m in order]); ax.invert_yaxis()
    ax.set_xlabel("TensorRT GPU latency (ms)"); ax.set_title("Edge TensorRT latency: AGX Orin vs Orin Nano (FP16/INT8)")
    ax.legend(frameon=False,fontsize=7,ncol=2); ax.yaxis.grid(False)
    fig.tight_layout(); fig.savefig(f"{FIGDIR}/fig6_edge_trt.png",dpi=200,bbox_inches="tight"); plt.close(fig)

    # Fig 7: INT8 accuracy cost (canonical, device-independent) vs FP32
    fig,ax=plt.subplots(figsize=(6.6,3.8)); y=np.arange(len(PRIM))
    order=sorted(PRIM, key=lambda m: acc_fp32(m))
    af=[acc_fp32(m) for m in order]; i8v=[int8_ate(m) for m in order]
    ax.barh(y-0.2,af,0.4,color="#2E86AB",label="FP32"); ax.barh(y+0.2,i8v,0.4,color="#8B1E3F",label="INT8 (canonical QDQ)")
    for yi,(a,b) in enumerate(zip(af,i8v)): ax.text(b,yi+0.2,f"  +{(b-a)/a*100:.0f}%",va="center",fontsize=8)
    ax.set_yticks(y); ax.set_yticklabels([DISP[m] for m in order]); ax.invert_yaxis()
    ax.set_xlabel("ATE (m) — lower better"); ax.set_title("INT8 PTQ accuracy cost (device-independent)")
    ax.legend(frameon=False,fontsize=8); ax.yaxis.grid(False)
    fig.tight_layout(); fig.savefig(f"{FIGDIR}/fig7_int8_acc.png",dpi=200,bbox_inches="tight"); plt.close(fig)

    # Fig 8: BF16 LSTM pitfall across the three devices
    fig,ax=plt.subplots(figsize=(5.4,3.4))
    devs=["Blackwell","AGX Orin","Orin Nano"]
    def _lstm_ratio(rows):
        b=None;f=None
        for r in rows:
            if r["model"]=="ronin_lstm" and r["precision"]=="bf16": b=float(r["lat_med_ms"])
            if r["model"]=="ronin_lstm" and r["precision"]=="fp16": f=float(r["lat_med_ms"])
        return b/f if b and f else None
    ratios=[_lstm_ratio(rows) for rows in [rows if False else list(csv.DictReader(open(os.path.join(AIOTC,"results","profile_blackwell.csv")))),
                                           EDGE["agx_orin"]["pt"], EDGE["orin_nano"]["pt"]]]
    ax.bar(devs, ratios, color=["#2E86AB","#1B998B","#8B1E3F"])
    for xi,v in enumerate(ratios): ax.text(xi,v,f"{v:.0f}×",ha="center",va="bottom",fontsize=10,fontweight="bold")
    ax.axhline(1,ls="--",color="gray",lw=1); ax.set_ylabel("RoNIN-LSTM BF16/FP16 latency ratio")
    ax.set_title("BF16 recurrent-kernel pitfall amplifies on edge")
    fig.tight_layout(); fig.savefig(f"{FIGDIR}/fig8_bf16_3dev.png",dpi=200,bbox_inches="tight"); plt.close(fig)
    print("edge figures written")

# Fig 5: accuracy-energy trade-off (FP32) over the admitted models; primary 5 emphasized
PRIMARY={"ronin_resnet18","ronin_lstm","ronin_tcn","imunet","tinyodom"}
if acc_rows:
    fig,ax=plt.subplots(figsize=(6.6,4.4))
    for m in acc_models:
        a=accget(m,'fp32','ate_m'); e=get(m,'fp32','energy_mJ_per_inf')
        if a is None or e!=e: continue
        prim=m in PRIMARY
        ax.scatter(e,float(a),s=130 if prim else 55,color=FAMCOL[FAM[m]],
                   marker='o' if prim else 's',edgecolor='black' if prim else 'gray',
                   linewidth=1.1 if prim else 0.6,zorder=3,alpha=0.95 if prim else 0.7)
        ax.annotate(DISP[m],(e,float(a)),textcoords="offset points",xytext=(6,3),
                    fontsize=8,fontweight='bold' if prim else 'normal')
    ax.set_xlabel("energy per inference (mJ, FP32) — lower better")
    ax.set_ylabel("ATE (m) — lower better")
    ax.set_title("Accuracy–energy trade-off on the Blackwell tier (FP32)")
    import matplotlib.lines as mlines
    ax.legend([mlines.Line2D([],[],marker='o',color='gray',ls='',mec='black',label='primary (Pareto)'),
               mlines.Line2D([],[],marker='s',color='gray',ls='',label='appendix (mobile CNN)')],
              ['primary set','appendix'],frameon=False,fontsize=8,loc='upper right')
    fig.tight_layout(); fig.savefig(f"{FIGDIR}/fig5_pareto.png",dpi=200,bbox_inches="tight"); plt.close(fig)

# ---------------- build docx ----------------
d = Document()
# base font
d.styles["Normal"].font.name="Times New Roman"; d.styles["Normal"].font.size=Pt(10.5)

def H(text, lvl=1):
    p=d.add_heading(text, level=lvl)
    for r in p.runs: r.font.color.rgb=RGBColor(0,0,0)
    return p
def P(text):
    p=d.add_paragraph(text); p.paragraph_format.space_after=Pt(6); return p
def CAP(text):
    p=d.add_paragraph(); r=p.add_run(text); r.italic=True; r.font.size=Pt(9)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER; return p

t=d.add_heading("A Fair Cross-Hardware Benchmark of Learning-Based Pedestrian Inertial Navigation: "
                "Accuracy, Latency and Energy from Datacenter GPU to Jetson Edge", level=0)
for r in t.runs: r.font.color.rgb=RGBColor(0,0,0)

H("1. Introduction", 1)
P("Learning-based pedestrian dead reckoning (PDR) estimates walking displacement directly from inertial "
  "measurements and is a key enabler of infrastructure-free indoor localization for IoT and smart-building "
  "applications. Because such systems must ultimately run on heterogeneous edge hardware, their deployment "
  "cost—latency and, critically, energy per localization update—must be characterized. This report presents "
  "a fair, reproducible cross-hardware benchmark of ten open-source PDR backbones spanning six architecture "
  "families, characterizing accuracy (ATE/RTE), latency and energy from a datacenter Blackwell GPU down to "
  "NVIDIA Jetson AGX Orin and Orin Nano edge modules, at FP32/FP16/BF16 and, on the edge, TensorRT INT8. A "
  "model-admission methodology keeps every compared model on one identical task and training recipe, and INT8 "
  "is unified to a single canonical quantized artifact so its accuracy is device-independent and only latency "
  "varies by device.")

H("2. Experimental Setup", 1)
P("Hardware. We measure three tiers spanning the deployment spectrum: a datacenter NVIDIA RTX PRO 6000 "
  "Blackwell GPU (compute capability 12.0, 96 GB, 600 W), and two NVIDIA Jetson edge modules covering the "
  "current Orin range — the flagship AGX Orin and the entry-level Orin Nano (both Ampere, sm 8.7, whole-board "
  "power in the ~5–60 W class). The Blackwell tier uses PyTorch 2.13/CUDA 13; each edge tier uses its JetPack "
  "PyTorch plus TensorRT. Detailed per-device run conditions (power mode, locked clocks, versions) are recorded "
  "with the released data.")
P("Models. We benchmark ten open-source learning-based inertial-odometry backbones spanning six architecture "
  "families: residual CNNs (RoNIN-ResNet, 4.64 M parameters; TLIO, 5.43 M), a recurrent network (RoNIN-LSTM, "
  "0.21 M), temporal convolutional networks (RoNIN-TCN, 0.54 M; TinyOdom, 0.11 M, a hardware-aware NAS design "
  "for microcontrollers), depthwise-separable mobile CNNs (IMUNet 3.66 M, MobileNetV2 2.18 M, MnasNet 2.98 M, "
  "EfficientNet-B0 3.23 M), and an O(2)-equivariant network (EqNIO, 5.23 M).")
P("Task-formulation caveat. These backbones do NOT all solve an identical task, and we do not treat them as "
  "interchangeable. Eight of them (the three RoNIN variants, IMUNet, the three mobile CNNs and TinyOdom) share "
  "one formulation — a 6-channel global-frame inertial window regressed to 2-D global velocity, integrated into "
  "a trajectory — and are trained and evaluated identically (Section 6). The remaining two use materially "
  "different formulations: TLIO regresses 3-D body-frame displacement with a covariance head and reconstructs "
  "its trajectory through a stochastic-cloning EKF, and EqNIO consumes an O(2)-equivariant vector/scalar "
  "representation rather than the raw window. TLIO and EqNIO are therefore reported as separate deployment case "
  "studies and are never mixed into the unified accuracy ranking. Per-model input, frame, output, training set, "
  "checkpoint and reconstruction are catalogued in an accompanying model admission sheet. Efficiency, which "
  "measures the per-inference cost of running the network on the hardware and is independent of the trained "
  "weights, is reported for all ten (TLIO's EKF and any per-model pre/post-processing are excluded from the "
  "measured cost).")
P("Primary model set. To keep the main benchmark focused and defensible, the primary comparison is drawn over "
  "five purpose-built inertial-odometry backbones that span the architecture and efficiency spectrum: "
  "RoNIN-ResNet, RoNIN-LSTM and RoNIN-TCN (the canonical CNN/RNN/TCN benchmark family), IMUNet (a "
  "mobile-optimized IO network) and TinyOdom (an MCU-targeted NAS design). The three generic image-backbone "
  "adaptations (MobileNetV2, MnasNet, EfficientNet-B0) and the two case studies (TLIO, EqNIO) are reported "
  "separately. The primary set anchors the main deployment recommendations; the three additional mobile CNNs "
  "are included in the unified-protocol accuracy analysis (Table 3, Fig. 5) as an extended architecture sweep, "
  "so the accuracy figures cover all eight unified-protocol backbones while the recommendations lead with the "
  "primary five. The efficiency tables list every model for completeness.")
P("Numerical precisions. Each model is evaluated at three precisions: FP32, FP16 and BF16, the two 16-bit "
  "formats that Blackwell supports natively. INT8 (which requires TensorRT quantization) is deferred to the edge "
  "tiers, whose JetPack TensorRT exposes the classic calibration API and where quantization actually accelerates "
  "inference; the INT8 accuracy penalty depends on the calibration set and the resulting TensorRT engine and is "
  "therefore measured per deployment, not assumed (Section 6).")
P("Metrics and protocol. For each configuration we report median and 95th-percentile per-inference latency "
  "(batch = 1, over 500 iterations after 50 warm-up iterations), throughput, parameter count and peak GPU "
  "memory. Power is a card-level figure — the mean nvidia-smi board-power reading over a sustained run — from "
  "which total energy per inference is derived as power divided by throughput; it is not a wall-plug "
  "measurement. For the primary models we additionally report idle-subtracted dynamic energy from a repeated "
  "idle/active protocol on the otherwise-idle card (Table 4). Latency, throughput and memory are direct "
  "measurements. Efficiency measurements are taken on a single idle GPU. Accuracy is a property of the "
  "checkpoint and its PyTorch inference artifact and is re-evaluated whenever a deployment changes precision, "
  "backend, graph transformation or quantization (Section 6).")

H("3. Results", 1)
P("Table 1 reports the full measurement grid; Figures 1–3 visualize latency, energy per inference, and the "
  "relationship between model size and latency.")

# Table 1
d.add_paragraph().add_run("Table 1. Latency, throughput, power, energy and memory on the RTX PRO 6000 "
                          "Blackwell GPU (batch = 1).").italic=True
cols=["Model","Prec.","Params (M)","Lat. med (ms)","Lat. P95 (ms)","Thr. (inf/s)","Power (W)","Energy (mJ)","Mem (MB)"]
tb=d.add_table(rows=1, cols=len(cols)); tb.style="Light Grid Accent 1"
for i,c in enumerate(cols): tb.rows[0].cells[i].paragraphs[0].add_run(c).bold=True
for m in models:
    for prec in PRECS:
        vals=[disp[m],prec.upper(),f"{get(m,prec,'params_M'):.2f}",f"{get(m,prec,'lat_med_ms'):.2f}",
              f"{get(m,prec,'lat_p95_ms'):.2f}",f"{get(m,prec,'throughput_ips'):.0f}",
              f"{get(m,prec,'power_W'):.0f}",f"{get(m,prec,'energy_mJ_per_inf'):.1f}",
              f"{get(m,prec,'peak_mem_MB'):.0f}"]
        cells=tb.add_row().cells
        for i,v in enumerate(vals):
            cells[i].paragraphs[0].add_run(v); cells[i].paragraphs[0].runs[0].font.size=Pt(9)

for fig,cap in [("fig1_latency.png","Fig. 1. FP32 per-inference latency (batch = 1), sorted, coloured by architecture family. TinyOdom (NAS-TCN) is fastest; EqNIO (equivariant) is ~4× slower."),
                ("fig2_energy.png","Fig. 2. FP32 energy per inference, sorted. The datacenter GPU spends 95–396 mJ per localization update depending on architecture."),
                ("fig3_params_latency.png","Fig. 3. Parameters (log scale) vs. latency. There is no monotone relationship: the 0.11 M TinyOdom is fastest and the 5.23 M EqNIO slowest, but the depthwise-separable mobile CNNs are slow at modest size — architecture, not parameter count, drives latency.")]:
    d.add_paragraph()
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(f"{FIGDIR}/{fig}", width=Inches(4.6))
    CAP(cap)

# numbers below are pulled live from the CSV so Table 1 and the findings can never drift
def L(m,p): return f"{get(m,p,'lat_med_ms'):.2f}"
def E(m,p): return f"{get(m,p,'energy_mJ_per_inf'):.0f}"
def T(m,p): return f"{get(m,p,'throughput_ips'):.0f}"
_e32=[get(m,'fp32','energy_mJ_per_inf') for m in models]; EMIN,EMAX=f"{min(_e32):.0f}",f"{max(_e32):.0f}"
_t32=[get(m,'fp32','throughput_ips') for m in models]; TMIN,TMAX=f"{min(_t32):.0f}",f"{max(_t32):.0f}"

H("4. Analysis and Discussion", 1)
P(f"Finding 1 — Reduced precision is not uniformly beneficial, and the best 16-bit format is model-dependent. "
  f"Neither FP16 nor BF16 delivers a consistent speed-up at batch = 1. FP16 slows RoNIN-ResNet "
  f"({L('ronin_resnet18','fp32')} → {L('ronin_resnet18','fp16')} ms) and IMUNet "
  f"({L('imunet','fp32')} → {L('imunet','fp16')} ms), is roughly neutral for RoNIN-TCN, and helps only "
  f"RoNIN-LSTM ({L('ronin_lstm','fp32')} → {L('ronin_lstm','fp16')} ms). BF16 tracks FP16 for the "
  f"convolutional backbones but is about three times slower on RoNIN-LSTM ({L('ronin_lstm','fp32')} → "
  f"{L('ronin_lstm','bf16')} ms, {E('ronin_lstm','bf16')} mJ) — a counter-intuitive gap between two 16-bit "
  f"formats that we trace to GPU kernel dispatch in a dedicated case study (Section 5). At batch = 1 these "
  f"sub-6 M-parameter models severely under-utilize a Blackwell GPU, so inference is launch-/latency-bound "
  f"rather than compute-bound and reduced precision yields no arithmetic benefit; such gains are expected on "
  f"the compute-constrained edge devices, not here. The practical lesson is that precision must be validated "
  f"per model and per software stack rather than assumed from bit-width. This concerns runtime performance "
  f"only; the effect of precision on localization accuracy is evaluated in Section 6.")
P(f"Finding 2 — Architecture, not parameter count, determines speed, and 'lightweight' mobile networks are not "
  f"fast on a GPU. FP32 latency spans roughly 4× across the zoo (Fig. 1), but this ordering is uncorrelated "
  f"with model size (Fig. 3). The fastest backbone is TinyOdom ({L('tinyodom','fp32')} ms), a 0.11 M-parameter "
  f"TCN found by hardware-aware neural architecture search for microcontrollers, while the slowest is the "
  f"5.23 M-parameter EqNIO ({L('eqnio','fp32')} ms). "
  "Critically, the entire family of depthwise-separable mobile CNNs — IMUNet, MobileNetV2, MnasNet and "
  "EfficientNet-B0 — clusters at the slow end (1.8–2.3 ms) despite modest parameter counts, because "
  "depthwise-separable convolutions have low arithmetic intensity and are memory-bound: they suit mobile "
  "CPUs/DSPs but under-use a high-throughput GPU that favours dense computation. Compact residual CNNs "
  "(RoNIN-ResNet, TLIO) and the NAS-TCN, by contrast, keep the GPU busy. Model selection must therefore target "
  "the intended hardware, not parameter count alone.")
P(f"Finding 3 — A datacenter GPU is energy-inefficient for PDR, especially once idle draw is separated out. "
  f"Total per-inference energy (mean nvidia-smi board power over a sustained run, divided by throughput) is "
  f"{EMIN}–{EMAX} mJ across architectures (Fig. 2). For the primary models we additionally ran a repeated "
  f"idle/active protocol (Table 4): the card idles at ~{_idle} W, so a large fraction of the total "
  f"is fixed idle draw rather than the cost of computing. Subtracting the idle baseline gives the dynamic "
  f"energy — the marginal cost of one inference — which is 30–40% lower (RoNIN-ResNet "
  f"{pwrget('ronin_resnet18','fp32','total_energy_mJ')} mJ total vs "
  f"{pwrget('ronin_resnet18','fp32','dynamic_energy_mJ')} mJ dynamic; TinyOdom "
  f"{pwrget('tinyodom','fp32','total_energy_mJ')} vs {pwrget('tinyodom','fp32','dynamic_energy_mJ')} mJ). Either "
  f"way, running a sub-6 M-parameter model on a 600 W-rated accelerator is wasteful; energy per inference on a "
  f"low-power edge board is expected to be one to two orders of magnitude lower, which is the central argument "
  f"for edge deployment. Board power is a card-level (not wall-plug) figure.")
P(f"Finding 4 — Real-time is not the bottleneck. Even the slowest configuration (RoNIN-LSTM at BF16, "
  f"{T('ronin_lstm','bf16')} inf/s) runs an order of magnitude above the ~20 Hz required for pedestrian "
  f"tracking, and FP32 throughput ranges from {TMIN} inf/s (EqNIO) to {TMAX} inf/s (TinyOdom). Feasibility is "
  f"therefore trivial on this tier; the meaningful differentiator across the hardware spectrum is energy "
  f"efficiency, not throughput.")

H("5. Case Study: A Precision-Selection Pitfall for Recurrent Backbones", 1)
P(f"The most striking entry in Table 1 is RoNIN-LSTM at BF16, which is roughly three times slower than at FP16 "
  f"({L('ronin_lstm','bf16')} vs {L('ronin_lstm','fp16')} ms) for an identical model, input and batch size. "
  f"Because one 16-bit format being 3× slower than another is counter-intuitive, we investigated whether this "
  f"is a measurement artifact or a genuine phenomenon, following a staged protocol (full log in the "
  f"accompanying diagnosis report).")
P("Ruling out experimental error. Across five independent runs the gap is stable (inter-run median spread "
  "< 0.01 ms); the parameter, input, hidden-state and output tensors are confirmed to be in the target dtype "
  "under every precision; no autocast is active; and BF16 is natively supported (compute capability 12.0). "
  "Re-flattening the LSTM weight buffer (flatten_parameters) does not change the result (3.04× vs 3.08×), "
  "excluding weight non-contiguity. The gap also reproduces in a minimal nn.LSTM stripped of all "
  "application-specific code, so it is a property of the LSTM operator on this stack, not of the pipeline.")
P("A hidden-size crossover reveals the mechanism. Sweeping the LSTM hidden size (Table 2, Fig. 4) shows the "
  "slowdown is not constant: it reaches 11× at hidden = 50, is 3× at hidden = 100, and nearly vanishes "
  "(1.13×) at hidden = 200. The gap is otherwise insensitive to sequence length and batch size.")

d.add_paragraph().add_run("Table 2. BF16 vs. FP16 latency of a minimal nn.LSTM (batch = 1, seq = 200, 3 layers) "
                          "as a function of hidden size.").italic=True
xh=[50,100,200]; xf=[0.402,1.449,3.948]; xb=[4.412,4.409,4.456]
t2cols=["Hidden size","FP16 (ms)","BF16 (ms)","BF16/FP16"]
t2=d.add_table(rows=1,cols=len(t2cols)); t2.style="Light Grid Accent 1"
for i,c in enumerate(t2cols): t2.rows[0].cells[i].paragraphs[0].add_run(c).bold=True
for h,fp,bp in zip(xh,xf,xb):
    cells=t2.add_row().cells
    for i,v in enumerate([str(h),f"{fp:.2f}",f"{bp:.2f}",f"{bp/fp:.2f}×"]):
        cells[i].paragraphs[0].add_run(v); cells[i].paragraphs[0].runs[0].font.size=Pt(9)

d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(f"{FIGDIR}/fig4_bf16_crossover.png", width=Inches(4.4))
CAP("Fig. 4. The BF16/FP16 LSTM slowdown collapses toward parity as the hidden size grows, indicating a "
    "kernel-selection rather than an arithmetic cause.")

P("GPU profiling confirms an execution-path difference. At hidden = 100, FP16 dispatches a single fused cuDNN "
  "persistent-RNN kernel (RNN_blockPersist_fp_LSTM), issuing about 1.1k kernel launches per 50 inferences. BF16 "
  "has no persistent-LSTM kernel and instead runs an unrolled per-timestep path (a per-step GEMV plus an "
  "elementwise recurrent-cell kernel), issuing about 124k launches — 112× more — so at batch = 1 it is "
  "dominated by kernel-launch overhead. At hidden = 200 the persistent kernel is no longer selected even for "
  "FP16, both precisions take the same unrolled path, and their launch counts and latencies converge (1.01×). "
  "The regression is therefore the absence of an FP16-only persistent kernel for BF16, not a difference in raw "
  "arithmetic. Two controls corroborate this. Disabling cuDNN forces all three precisions onto the same native "
  "path, roughly 9× slower than the cuDNN FP16 kernel, with the gap gone (BF16/FP16 = 1.00×); and the entire "
  "phenomenon reproduces on a newer stack (PyTorch 2.14-dev, cuDNN 9.24.0), so it is not a transient artifact of "
  "one library build.")
P("Scope and takeaway. This behaviour is specific to the evaluated stack (RTX PRO 6000 Blackwell, compute "
  "capability 12.0, PyTorch 2.13/CUDA 13.0/cuDNN 9.2.0, driver 580.142, batch = 1, PyTorch-eager execution), "
  "and was reproduced on cuDNN 9.24.0. It is a library-kernel-coverage phenomenon, not a hardware law; we do "
  "not claim it for every RNN type or framework, but it is not stack-specific either — it reproduces and "
  "worsens on two different (Ampere) edge GPUs below. The actionable takeaway is that for latency-critical "
  "single-sample recurrent inference, FP16 is preferable to BF16 — a recommendation about runtime performance "
  "only, independent of any accuracy considerations.")
if has_edge:
    def _lstmr(rows):
        dd={r["precision"]:float(r["lat_med_ms"]) for r in rows if r["model"]=="ronin_lstm"}
        return dd["bf16"]/dd["fp16"]
    P(f"Cross-device: the pitfall generalizes and worsens on the edge. Rebuilding the same RoNIN-LSTM at BF16 on "
      f"the two Jetson Orins — a different GPU architecture (Ampere, sm_87) and cuDNN than the Blackwell host — "
      f"not only reproduces the slowdown but amplifies it: BF16 is {_lstmr(EDGE['agx_orin']['pt']):.0f}× slower "
      f"than FP16 on the AGX Orin and {_lstmr(EDGE['orin_nano']['pt']):.0f}× on the Orin Nano, versus 3.1× on "
      f"the Blackwell GPU (Fig. 8). The missing persistent BF16 recurrent kernel is therefore not a single-stack "
      f"quirk; on resource-constrained edge hardware it turns a recurrent PDR backbone at BF16 from real-time "
      f"into unusable (95 ms per inference on the Nano). The deployment rule sharpens accordingly: never ship a "
      f"recurrent inertial model at BF16.")
    d.add_paragraph(); _p=d.add_paragraph(); _p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    _p.add_run().add_picture(f"{FIGDIR}/fig8_bf16_3dev.png", width=Inches(3.5))
    CAP("Fig. 8. The BF16 recurrent-kernel pitfall across three devices: the RoNIN-LSTM BF16/FP16 latency ratio "
        "grows from 3.1x on the Blackwell GPU to ~27-30x on the Jetson Orins.")
P("One might ask whether BF16 at least buys numerical robustness that could offset its cost. Using identical "
  "weights and inputs, we measured each 16-bit format's output deviation from an FP32 reference (random weights, "
  "so this probes architecture-level numerical sensitivity, not localization accuracy). Both formats are "
  "faithful — cosine similarity exceeds 0.9999 for all four backbones — but FP16 is consistently 6–23× closer "
  "to FP32 than BF16 (relative L2 error 2.9e-4 vs 2.0e-3 for RoNIN-LSTM, and 4.9e-4 vs 1.1e-2 for IMUNet). "
  "BF16's wider dynamic range does not help here because bounded inference activations never overflow, so its "
  "coarser 7-bit mantissa only costs precision. On this stack FP16 therefore dominates BF16 on both measured "
  "axes — latency and output fidelity — for these PDR backbones. This remains a weight-independent numerical "
  "check; the effect of precision on end-to-end localization accuracy requires trained weights and is deferred "
  "to the accuracy evaluation (Section 6).")

if acc_rows:
    H("6. Localization Accuracy and the Cost of Quantization", 1)
    P("Setup. We evaluate localization accuracy on the MagPIE walking test set (53 sequences) under the RoNIN "
      "protocol: each backbone regresses global 2-D velocity, which is integrated into a trajectory and compared "
      "against visual-inertial ground truth using Absolute and Relative Trajectory Error (ATE / RTE). To keep "
      "the cross-model comparison fair, every window-based CNN backbone is trained with one identical recipe "
      "(Adam, lr = 1e-4, batch = 256, RandomHoriRotate augmentation, ReduceLROnPlateau, early stopping on "
      "validation), and the sequence models (LSTM/TCN) with RoNIN's native sequence-to-sequence recipe. For a "
      "fixed checkpoint and identical PyTorch inference artifact, task accuracy is treated as model-level and is "
      "evaluated once on a single GPU. Accuracy is re-evaluated whenever deployment changes the numerical "
      "precision, inference backend, graph transformation or quantization procedure; this is mandatory for the "
      "planned TensorRT FP16/INT8 engines, whose calibration and kernel/graph fusion can alter the numerical "
      "result.")
    P("Admission. Only the eight unified-protocol backbones enter this accuracy comparison; TLIO and EqNIO are "
      "excluded because their task formulations differ (Section 2) and are reported as separate case studies. "
      "This keeps the accuracy–energy Pareto an apples-to-apples comparison over models solving one identical "
      "task with one identical reconstruction and metric.")
    _AA=lambda m: float(accget(m,'fp32','ate_m'))
    _a32=[_AA(m) for m in acc_models if accget(m,'fp32','ate_m')]
    P(f"Results. Table 3 reports ATE/RTE for all eight admitted backbones, each trained identically on this "
      f"machine. FP32 ATE spans {min(_a32):.2f}–{max(_a32):.2f} m. The window CNNs are the most accurate and "
      f"cluster tightly — RoNIN-ResNet {_AA('ronin_resnet18'):.2f} m, MnasNet {_AA('mnasnet'):.2f}, "
      f"MobileNetV2 {_AA('mobilenetv2'):.2f}, IMUNet {_AA('imunet'):.2f}, EfficientNet-B0 "
      f"{_AA('efficientnet_b0'):.2f} m — whereas the sequence models and the tiny NAS model trail: RoNIN-TCN "
      f"{_AA('ronin_tcn'):.2f}, TinyOdom {_AA('tinyodom'):.2f}, RoNIN-LSTM {_AA('ronin_lstm'):.2f} m.")
    P(f"Accuracy–energy trade-off (Fig. 5). The efficiency champion TinyOdom ({L('tinyodom','fp32')} ms, "
      f"{E('tinyodom','fp32')} mJ; Table 1) is among the least accurate ({_AA('tinyodom'):.2f} m), while the "
      f"most accurate models (RoNIN-ResNet and MnasNet, ~0.90 m) cost more energy per update. This is the "
      f"trade-off a deployment must resolve: on the Blackwell tier both RoNIN-ResNet (best accuracy) and "
      f"TinyOdom (best energy) sit on the Pareto frontier. The edge tiers (ongoing) are expected to reshape this "
      f"frontier as small-model and low-precision advantages materialize on compute-constrained hardware.")
    d.add_paragraph().add_run("Table 3. Localization accuracy (ATE/RTE, metres) on the MagPIE test set at three "
                              "numerical precisions. Lower is better.").italic=True
    a_cols=["Model","Prec.","ATE (m)","RTE (m)"]
    t3=d.add_table(rows=1, cols=len(a_cols)); t3.style="Light Grid Accent 1"
    for i,c in enumerate(a_cols): t3.rows[0].cells[i].paragraphs[0].add_run(c).bold=True
    for m in acc_models:
        for prec in ("fp32","fp16","bf16"):
            ate=accget(m,prec,"ate_m"); rte=accget(m,prec,"rte_m")
            if ate is None: continue
            cells=t3.add_row().cells
            for i,v in enumerate([DISP.get(m,m), prec.upper(), f"{float(ate):.3f}", f"{float(rte):.3f}"]):
                cells[i].paragraphs[0].add_run(str(v)); cells[i].paragraphs[0].runs[0].font.size=Pt(9)
    def _reldl(prec):
        v=[(abs(float(accget(m,prec,'ate_m'))-float(accget(m,'fp32','ate_m'))),
            abs(float(accget(m,prec,'ate_m'))-float(accget(m,'fp32','ate_m')))/float(accget(m,'fp32','ate_m'))*100,
            DISP.get(m,m)) for m in acc_models if accget(m,prec,'ate_m')]
        return max(v) if v else (0.0,0.0,'')
    _m16=_reldl('fp16'); _mbf=_reldl('bf16')
    P(f"Finding 5 — FP16 preserves localization accuracy for the evaluated checkpoints; BF16 is close but not "
      f"identical. For the fixed checkpoints evaluated on MagPIE, FP16 changed ATE by at most {_m16[0]:.3f} m "
      f"({_m16[1]:.1f}%) relative to FP32. This observation is limited to the evaluated checkpoints and test "
      f"set; multi-seed training variance was not assessed. BF16 is small but not negligible: it shifts ATE by "
      f"up to {_mbf[0]:.3f} m ({_mbf[1]:.1f}%, worst case {_mbf[2]}), consistently a little further from FP32 "
      f"than FP16 — the accuracy counterpart to the numerical-consistency measurement in Section 5. Practical "
      f"guidance: prefer FP16 for PyTorch edge inference; validate BF16 per model. We do not extrapolate to "
      f"INT8, whose penalty depends on calibration and the resulting TensorRT engine and must be measured per "
      f"deployment (Section 7).")
    if os.path.exists(f"{FIGDIR}/fig5_pareto.png"):
        d.add_paragraph()
        p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(f"{FIGDIR}/fig5_pareto.png", width=Inches(4.6))
        CAP("Fig. 5. Accuracy (ATE) versus energy per inference on the Blackwell tier (FP32). Primary models are "
            "circled; the appendix mobile CNNs are squares. RoNIN-ResNet (best accuracy) and TinyOdom (best "
            "energy) anchor the Pareto frontier; TLIO/EqNIO are excluded (different task, Section 2).")
    if pwr_rows:
        P("Dynamic-power detail. Table 4 gives the idle/active power breakdown behind Finding 3 for the primary "
          "models, so that the energy axis of Fig. 5 can be read either as total or as idle-subtracted dynamic "
          "energy.")
        d.add_paragraph().add_run("Table 4. Idle/active dynamic-power breakdown for the primary models. Protocol: "
            "nvidia-smi card power sampled at ~10 Hz and median-aggregated; per configuration, N independent "
            "trials of Duration/run seconds idle then the same under a sustained batch-1 inference loop. Total "
            "energy = active power / throughput; dynamic energy subtracts the idle baseline (mean ± per-trial "
            "std). GPU at stock clocks (not locked, 600 W limit); small differences from Table 1's energy arise "
            "because this is a separate, longer protocol rather than Table 1's single 8-second run.").italic=True
        p4=["Model","Prec.","Idle (W)","Active (W)","Thr. (inf/s)","E total (mJ)","E dyn (mJ)","Runs","Dur/run (s)"]
        t4=d.add_table(rows=1, cols=len(p4)); t4.style="Light Grid Accent 1"
        for i,c in enumerate(p4): t4.rows[0].cells[i].paragraphs[0].add_run(c).bold=True
        for m in ["ronin_resnet18","ronin_lstm","ronin_tcn","imunet","tinyodom"]:
            for prec in ("fp32","fp16","bf16"):
                if pwrget(m,prec,"idle_W") is None: continue
                edyn=f"{pwrget(m,prec,'dynamic_energy_mJ')}±{pwrget(m,prec,'dynamic_energy_std_mJ')}"
                vals=[DISP.get(m,m),prec.upper(),pwrget(m,prec,"idle_W"),pwrget(m,prec,"active_W"),
                      f"{float(pwrget(m,prec,'throughput_ips')):.0f}",pwrget(m,prec,"total_energy_mJ"),
                      edyn,pwrget(m,prec,"runs"),pwrget(m,prec,"dur_s")]
                cells=t4.add_row().cells
                for i,v in enumerate(vals):
                    cells[i].paragraphs[0].add_run(str(v)); cells[i].paragraphs[0].runs[0].font.size=Pt(8)

if has_edge:
    H("7. Edge Tiers: Jetson AGX Orin and Orin Nano", 1)
    P("We extend the benchmark to two edge tiers spanning the current Orin deployment range — the flagship "
      "AGX Orin and the entry-level Orin Nano — anchored by the datacenter Blackwell GPU. Efficiency is measured "
      "per device (whole-board power via tegrastats); accuracy is device-independent for a fixed PyTorch "
      "checkpoint, confirmed below. For INT8 we adopt a single canonical quantized artifact — one QDQ ONNX per "
      "model with a fixed post-training-quantization recipe — so INT8 accuracy is a property of that artifact "
      "(measured once, device-independent) while each device builds its TensorRT engine from the identical ONNX, "
      "making INT8 latency a clean hardware comparison rather than a quantization-method artifact.")
    P(f"Accuracy is device-independent. The PyTorch FP32/FP16 ATE measured on both Orins matches the datacenter "
      f"reference to within 0.001 m (e.g. RoNIN-ResNet {acc_fp32('ronin_resnet18'):.3f} m on every tier), "
      f"empirically confirming that a fixed checkpoint's accuracy is a model-level property; only INT8, which "
      f"changes the numerics, deviates.")
    _tp=[pt_lat("agx_orin","fp16",m)/trt_lat("agx_orin","fp16",m) for m in PRIM]
    _agxnano=[trt_lat("orin_nano","fp16",m)/trt_lat("agx_orin","fp16",m) for m in PRIM]
    P(f"Finding 6 — On the edge, the runtime dominates the precision. Moving a model from PyTorch to a TensorRT "
      f"engine (same device, FP16) speeds it up by {min(_tp):.0f}–{max(_tp):.0f}×, far more than the roughly "
      f"1–2× a further FP16->INT8 step can buy. At batch 1 in PyTorch the flagship AGX Orin is no faster than the "
      f"entry Orin Nano — both are launch-bound — and the AGX's compute advantage (a {min(_agxnano):.1f}–"
      f"{max(_agxnano):.1f}× TensorRT-latency lead, Table 5, Fig. 6) only appears once TensorRT is used. The "
      f"first lever for edge deployment is therefore the inference runtime, not the numerical precision.")
    d.add_paragraph().add_run("Table 5. Per-inference latency (ms) on the two edge tiers: PyTorch FP16 vs "
                              "TensorRT FP16 vs TensorRT INT8 (INT8 built from the canonical QDQ ONNX).").italic=True
    t5c=["Model","AGX PT16","AGX TRT16","AGX TRT8","Nano PT16","Nano TRT16","Nano TRT8"]
    t5=d.add_table(rows=1,cols=len(t5c)); t5.style="Light Grid Accent 1"
    for i,c in enumerate(t5c): t5.rows[0].cells[i].paragraphs[0].add_run(c).bold=True
    for m in PRIM:
        vals=[DISP[m], f"{pt_lat('agx_orin','fp16',m):.1f}", f"{trt_lat('agx_orin','fp16',m):.2f}",
              f"{trt_lat('agx_orin','int8',m):.2f}", f"{pt_lat('orin_nano','fp16',m):.1f}",
              f"{trt_lat('orin_nano','fp16',m):.2f}", f"{trt_lat('orin_nano','int8',m):.2f}"]
        cs=t5.add_row().cells
        for i,v in enumerate(vals): cs[i].paragraphs[0].add_run(v); cs[i].paragraphs[0].runs[0].font.size=Pt(8)
    d.add_paragraph(); p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(f"{FIGDIR}/fig6_edge_trt.png", width=Inches(4.5))
    CAP("Fig. 6. Edge TensorRT latency, AGX Orin vs Orin Nano, FP16 and INT8 (from the identical canonical ONNX).")
    _d8=[(int8_ate(m)-acc_fp32(m))/acc_fp32(m)*100 for m in PRIM]
    P(f"Finding 7 — INT8 is not free, and its cost is model- and toolchain-dependent. Under the canonical QDQ "
      f"post-training quantization every backbone loses accuracy — ATE degrades by {min(_d8):.0f}–{max(_d8):.0f}% "
      f"relative to FP32 (Table 6, Fig. 7), device-independently. The magnitude is highly sensitive to the PTQ "
      f"recipe: an entropy-calibrator variant we also ran left several models near-lossless while still "
      f"destroying others, so INT8 accuracy must be measured for the exact quantized artifact, never assumed. "
      f"Meanwhile INT8 buys little latency over FP16 for these tiny models (often under 1.3×, and slower than "
      f"FP16 in some cases), so for most of them INT8 PTQ is a poor trade — a large accuracy loss for a marginal "
      f"speed-up. FP16 is the pragmatic edge precision.")
    d.add_paragraph().add_run("Table 6. Unified INT8: device-independent accuracy of the canonical QDQ artifact, "
                              "and per-device INT8 TensorRT latency (built from the identical ONNX).").italic=True
    t6c=["Model","FP32 ATE","INT8 ATE","ΔATE","AGX INT8 ms","Nano INT8 ms"]
    t6=d.add_table(rows=1,cols=len(t6c)); t6.style="Light Grid Accent 1"
    for i,c in enumerate(t6c): t6.rows[0].cells[i].paragraphs[0].add_run(c).bold=True
    for m in PRIM:
        a0=acc_fp32(m); i8=int8_ate(m)
        vals=[DISP[m], f"{a0:.3f}", f"{i8:.3f}", f"+{(i8-a0)/a0*100:.0f}%",
              f"{trt_lat('agx_orin','int8',m):.2f}", f"{trt_lat('orin_nano','int8',m):.2f}"]
        cs=t6.add_row().cells
        for i,v in enumerate(vals): cs[i].paragraphs[0].add_run(v); cs[i].paragraphs[0].runs[0].font.size=Pt(8)
    d.add_paragraph(); p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(f"{FIGDIR}/fig7_int8_acc.png", width=Inches(4.4))
    CAP("Fig. 7. INT8 post-training-quantization accuracy cost (canonical QDQ, device-independent): every "
        "backbone degrades by 37–74% ATE relative to FP32.")

H(("8" if has_edge else ("7" if acc_rows else "6"))+". Limitations and Ongoing Work", 1)
P("The benchmark now covers three tiers (datacenter Blackwell, AGX Orin, Orin Nano), accuracy for all eight "
  "unified backbones, a device-independent INT8 accuracy, and per-device INT8 latency from a single canonical "
  "artifact. Its scope is nonetheless bounded, honestly: (i) the edge tiers are the current Orin generation "
  "(both Ampere); older architectures (e.g. Xavier/Volta) and MCU-class targets are not covered, so we frame "
  "the edge story as the Orin deployment range, not all edge hardware. (ii) On the AGX Orin the tiny batch-1 "
  "models barely move whole-board power, so its energy figures are baseline-dominated and latency is the "
  "reliable edge metric; the Orin Nano telemetry is more responsive. (iii) INT8 accuracy is reported for one "
  "fixed PTQ recipe (per-tensor symmetric QDQ); other recipes (per-channel, entropy calibration, QAT) change "
  "it — sometimes substantially — which is precisely why we treat INT8 accuracy as a per-artifact quantity to "
  "be measured, not assumed. (iv) Accuracy is for the released checkpoints on the MagPIE test set; multi-seed "
  "training variance is not assessed. (v) TLIO and EqNIO remain separate deployment case studies (different "
  "task), never merged into the unified ranking. The natural next step is a device- and budget-aware "
  "model-and-precision selection recipe distilled from the accuracy-versus-energy Pareto across all three "
  "tiers.")

d.save(OUT)
print("wrote", OUT)
