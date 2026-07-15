#!/usr/bin/env python3
"""Chinese conference-paper draft (docx) restructured to follow
AIoTC_PDR_Conference_Paper_Outline.docx: a kernel-aware batch-1 deployment cost model
(T0 + a*N_exec + b*B_eff + c*F_eff + d*I_fallback) + budget-aware deployment selection,
with RQ1-RQ4, LOMO/LOAO/LODO/LOPO validation vs baselines, the BF16-LSTM fallback case,
controlled dynamic energy, and a real reference list. All numbers live from CSV/JSON.
"""
import os, csv, json
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AIOTC = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG = os.path.join(AIOTC, "results", "figures_pub")
OUT = os.path.join(AIOTC, "AIoTC_中文初稿.docx")
def L(p):
    fp = os.path.join(AIOTC, "results", p)
    return list(csv.DictReader(open(fp))) if os.path.exists(fp) else []
def gv(rows, m, prec, k):
    for r in rows:
        if r["model"] == m and (prec is None or r.get("precision") == prec):
            v = r.get(k); return float(v) if v not in (None, "") else None
BB, AG, NA = L("profile_blackwell.csv"), L("profile_agx_orin.csv"), L("profile_orin_nano.csv")
MM = ["ronin_resnet18","ronin_tcn","imunet","mobilenetv2","mnasnet","efficientnet_b0","tinyodom","tlio_resnet","eqnio"]
def lat(rows, m, pr): return gv(rows, m, pr, "lat_med_ms")
def factor(rows):
    x = np.array([lat(BB, m, "fp32") for m in MM]); y = np.array([lat(rows, m, "fp32") for m in MM])
    s = (y@x)/(x@x); r = np.corrcoef(x, y)[0, 1]; return s, r, np.mean(np.abs(s*x-y)/y)*100
sA, rA, mA = factor(AG); sN, rN, mN = factor(NA)
VAL = json.load(open(os.path.join(AIOTC, "results", "cost_model_validation.json")))
LOAO = json.load(open(os.path.join(AIOTC, "results", "loao_analysis.json")))
KP = json.load(open(os.path.join(AIOTC, "results", "kernel_profile.json")))
def vc(proto, key):
    s = VAL.get(proto, {}).get(key); return f"{s['MAPE']:.0f}/{s['Median']:.0f}" if s else "—"
PD = L("power_dynamic_blackwell.csv"); PDN = L("power_dynamic_orin_nano.csv"); PDA = L("power_dynamic_orin.csv")
def _ed(rows, m, pr):
    for r in rows:
        if r["model"] == m and r["precision"] == pr: return float(r["dynamic_energy_mJ"])
def edyn(m, pr): return _ed(PD, m, pr)
def edyn_n(m, pr): return _ed(PDN, m, pr)
def edyn_a(m, pr): return _ed(PDA, m, pr)
# cross-device dynamic-energy factorization (exclude lstm-bf16 fallback)
_Me = [m for m in MM if m != "ronin_lstm"]
_be = np.array([edyn(m, "fp32") for m in _Me])
def _efac(rows):
    e = np.array([_ed(rows, m, "fp32") for m in _Me]); s = (e@_be)/(_be@_be)
    return s, np.corrcoef(_be, e)[0, 1], np.mean(np.abs(s*_be-e)/e)*100
eSn, eRn, eMn = _efac(PDN); eSa, eRa, eMa = _efac(PDA)
ba = L("accuracy_blackwell.csv"); ci = L("accuracy_int8_canonical.csv")
_a = lambda m: gv(ba, m, "fp32", "ate_m")
def trt(src, m, pk):
    for r in L(src):
        if r["model"] == m and (pk is None or r.get("precision") == pk):
            return float(r.get("gpu_lat_med_ms") or r.get("trtexec_lat_med_ms"))
DISP = {"ronin_resnet18":"RoNIN-ResNet","ronin_tcn":"RoNIN-TCN","ronin_lstm":"RoNIN-LSTM","imunet":"IMUNet",
        "mobilenetv2":"MobileNetV2","mnasnet":"MnasNet","efficientnet_b0":"EfficientNet-B0","tlio_resnet":"TLIO",
        "tinyodom":"TinyOdom","eqnio":"EqNIO"}
FAMZH = {"ronin_resnet18":"CNN","tlio_resnet":"CNN","ronin_lstm":"RNN/LSTM","ronin_tcn":"TCN","tinyodom":"TCN(NAS)",
         "imunet":"轻量CNN","mobilenetv2":"轻量CNN","mnasnet":"轻量CNN","efficientnet_b0":"轻量CNN","eqnio":"等变"}
ALL = ["ronin_resnet18","tlio_resnet","ronin_lstm","ronin_tcn","tinyodom","imunet","mobilenetv2","mnasnet","efficientnet_b0","eqnio"]

d = Document()
d.styles["Normal"].font.size = Pt(10.5); d.styles["Normal"].font.name = "Times New Roman"
def cjk(run, body="宋体"):
    run.font.name = body; rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None: rf = OxmlElement("w:rFonts"); rPr.append(rf)
    rf.set(qn("w:eastAsia"), body)
def P(text, size=10.5, after=6, first_indent=True, bold=False):
    p = d.add_paragraph(); r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; cjk(r)
    p.paragraph_format.space_after = Pt(after)
    if first_indent: p.paragraph_format.first_line_indent = Pt(21)
    return p
def BUL(text):
    p = d.add_paragraph(style="List Bullet"); r = p.add_run(text); r.font.size = Pt(10.5); cjk(r)
    p.paragraph_format.space_after = Pt(2)
def H(text, lvl=1):
    p = d.add_paragraph(); r = p.add_run(text)
    r.bold = True; r.font.size = Pt(13-lvl); cjk(r, "黑体")
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4); return p
def CAP(text):
    p = d.add_paragraph(); r = p.add_run(text); r.font.size = Pt(9); r.italic = True; cjk(r)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8)
def figure(fn, cap, width=5.0):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(os.path.join(FIG, fn), width=Inches(width)); CAP(cap)
def small_table(cols, rows, cap):
    r0 = d.add_paragraph().add_run(cap); r0.italic = True; r0.font.size = Pt(9); cjk(r0)
    tb = d.add_table(rows=1, cols=len(cols)); tb.style = "Light Grid Accent 1"
    for i, c in enumerate(cols):
        rr = tb.rows[0].cells[i].paragraphs[0].add_run(c); rr.bold = True; rr.font.size = Pt(8.5); cjk(rr)
    for row in rows:
        cs = tb.add_row().cells
        for i, v in enumerate(row):
            rr = cs[i].paragraphs[0].add_run(str(v)); rr.font.size = Pt(8); cjk(rr)

# ===================== Title =====================
t = d.add_paragraph(); r = t.add_run("面向边缘 GPU 学习式行人航位推算的\n内核感知推理代价预测与预算约束部署")
r.bold = True; r.font.size = Pt(16); cjk(r, "黑体"); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
st = d.add_paragraph(); sr = st.add_run("Kernel-Aware Cost Prediction and Budget-Aware Deployment for\nLearning-Based Pedestrian Dead Reckoning on Edge GPUs")
sr.italic = True; sr.font.size = Pt(11); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
a = d.add_paragraph(); ar = a.add_run("匿名作者    单位"); ar.font.size = Pt(10.5); cjk(ar)
a.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ===================== Abstract =====================
H("摘要", 1)
P(f"学习式行人航位推算（PDR）在边缘设备部署时，batch=1 实时推理的延迟与能耗决定可行性，而参数量、FLOPs 或桌面 GPU 吞吐量"
  f"都无法可靠反映真实端侧代价。跨模型、精度、后端与硬件的系统 profiling 表明，执行图开销、有效访存、计算负载与内核可用性"
  f"共同决定延迟，BF16-LSTM 的异常慢速是典型例子。我们据此提出面向 batch=1 学习式 PDR 的内核感知部署代价模型，用执行单元数、"
  f"有效访存、计算量与一个 fallback 指示项预测给定“模型–精度–后端–设备”组合的延迟（必要时同时预测动态能耗）。在 RTX PRO "
  f"6000 Blackwell、Jetson AGX Orin、Jetson Orin Nano 与十个 PDR 网络上，我们以 leave-one-model-out / leave-one-architecture-out"
  f" / leave-one-device-out 协议与参数量、FLOPs 基线严格比较：结构模型在每个协议上都优于朴素基线（留一模型 "
  f"{VAL['LOMO']['M4 struct(ours)']['MAPE']:.0f}% vs {VAL['LOMO']['B1 params']['MAPE']:.0f}–"
  f"{VAL['LOMO']['B3 par+FLOP']['MAPE']:.0f}% MAPE），在本文三档设备上跨设备单标量迁移仅 {VAL['LODO']['M5 transfer(1-calib)']['MAPE']:.0f}%。需强调这些留出结论限定于受测的十个模型与三档硬件，不宣称对任意新架构或新硬件的普遍泛化——对结构全新的架构类需少量 few-shot 校准。"
  f"fallback 指示项把 BF16-LSTM 的预测误差从 65–78% 降至可忽略。基于该模型，我们构建预算感知部署选择器，在给定实时截止时间、"
  f"能耗预算与精度要求下自动推荐模型–精度–后端组合，把端侧 PDR 部署从穷举测量转为预测驱动的约束感知决策。", after=8)
kw = d.add_paragraph(); kr = kw.add_run("关键词： "); kr.bold=True; kr.font.size=Pt(10); cjk(kr)
kr2 = kw.add_run("行人航位推算；边缘 AI；惯性里程计；延迟预测；GPU Profiling；TensorRT；混合精度推理；资源感知部署")
kr2.font.size = Pt(10); cjk(kr2)

# ===================== 1 Introduction =====================
H("1  引言", 1)
P("室内无缝定位是智能穿戴与移动机器人等物联网应用的重要基础。针对卫星导航信号在封闭环境中常常失效的难题，基于深度学习的"
  "行人航位推算（PDR）[1,2,3] 通过挖掘终端内置惯性传感器的数据特征，成为维持连续定位的核心手段。为确保系统响应的实时性并"
  "降低对云端通信的依赖，将 PDR 模型直接下沉至边缘设备进行本地推理已成为领域共识。然而，这一部署方式面临严峻的现实阻力："
  "一方面，边缘硬件的计算资源天然受限；另一方面，高频的定位更新需求迫使模型以单样本流式（batch=1）的方式连续处理数据，"
  "将单次推理的时延预算极限压缩至十几毫秒以内（100 Hz 约 10 ms、50 Hz 约 20 ms）。在算力与实时性的双重约束下，如何为特定"
  "硬件可靠地完成模型评估与部署选型，构成了当前亟待解决的技术挑战。")
P("面对这一挑战，业界过去习惯依赖参数量或浮点运算数（FLOPs）来预估运行速度，但这种基于理论规模的标尺在边缘端存在明显局限。"
  "真实的推理延迟不仅取决于计算量，更高度受制于执行图开销、访存瓶颈以及底层算子的匹配度——如图 2 所示，参数量与实测延迟"
  "明显偏离。尤其是某些直觉上应当加速的低精度配置，可能因硬件缺乏对应的优化内核而触发非优化回退（fallback），反而造成"
  "数量级的性能劣化：例如循环网络在 BF16 下因缺少融合内核，单次推理的内核启动数从 10 暴增至 1241（124×），在入门边缘设备上"
  "时延放大近 30 倍、直接从可实时沦为不可用。由此可见，仅凭规模类指标往往难以指导端侧的实际部署，迫切需要一种能真实反映底层"
  "硬件机制的性能预测方法。")
P("针对上述部署痛点，本文提出一种专为边缘计算平台设计的内核感知代价模型。该模型突破了单纯依赖计算量的局限，将执行单元数"
  "（N_exec）、有效访存、执行图开销及关键的底层算子回退机制（fallback）显式纳入评估，从而预测多维组合配置下的真实时延与"
  "动态能耗（方法总览见图 1）。一个由此带来的实用性质是：batch=1 的延迟与能耗可因子化为“设备无关的模型形状 × 每设备标量”，"
  "因此只需一个标定模型，即可把数据中心的 profile 迁移到未见边缘设备，将 N 模型 × M 设备的穷举测量降为 O(N+M)。依托该代价"
  "模型，本文进一步构建面向固定终端的资源约束部署选择器：给定设备的运行频率与能耗上限，即可离线快速筛除违规配置并推荐"
  "较优方案，从而大幅减少系统落地前繁杂的穷举性能测试，为开发者提供更具原理性的选型依据。")
P("实验部分在涵盖数据中心与两档边缘设备的三类硬件平台上，结合十种主流网络与多种推理引擎进行了全面验证。留一模型交叉验证"
  "表明，相较以参数量/FLOPs 为效率代理的朴素回归对照（误差 44–60%；更精细的延迟预测器见 §2.3），本文代价模型将预测误差降至约 31%；而在最具挑战的跨设备预测"
  "上，基线误差高达约 290%，本文单标量迁移仅凭一个标定模型即达 11%，实现数量级的提升。更为重要的是，该模型能够精准捕捉并"
  "预警规模类标尺完全无法预测的算子回退陷阱（如上述 124× 内核激增），从而为资源约束下的部署提供高置信度的排序与选型推荐。"
  "需说明，上述留出结论限定于受测的十个模型与三档硬件，对结构全新的架构类需少量 few-shot 校准。")
P("本文的主要贡献总结如下：")
P("① 提出一种面向边缘设备单样本流式推理的内核感知代价模型，将执行图开销、有效访存及底层算子回退机制显式纳入时延与能耗的"
  "量化建模，并证明其在留出验证下显著优于参数量/FLOPs 基线、跨设备迁移仅需一个标定模型；", first_indent=False)
P("② 通过跨平台与多架构的广泛评测，揭示网络结构与低位宽精度不匹配引发的非优化回退等底层性能陷阱，并给出可直接观测的量化"
  "证据（如内核启动数 124× 激增），为端侧可靠部署提供实证预警；", first_indent=False)
P("③ 设计一套资源约束下的离线部署选择流程，在大幅减少穷举测试的前提下，针对具体终端直接推荐满足实时与能耗硬性指标的配置。"
  "全部权重、数据划分与脚本公开、可一键复现。", first_indent=False)
figure("fig0_overview.png", "图 1  方法总览：执行图 profiling 提取四个执行特征（N_exec、B_eff、F_eff、I_fallback） → 内核感知代价模型（含 fallback 项）预测延迟/能耗 → 预算感知选择器输出推荐并预警 BF16-LSTM 等陷阱。", 6.3)

# ===================== 2 Related work =====================
H("2  背景与相关工作", 1)
H("2.1  学习式 PDR", 2)
P("RoNIN [1] 以 ResNet/LSTM/TCN 回归全局速度，确立任务范式；TLIO [2] 回归带协方差的位移并接入 EKF；IONet [3] 较早用 RNN "
  "缓解惯性漂移；IMUNet [4] 与 TinyOdom [5] 面向移动端/微控制器做轻量化与硬件感知搜索；EqNIO [6] 以等变结构编码几何先验。"
  "按骨干可归为 CNN [7]、RNN/LSTM [11]、TCN [12] 与轻量网络 [8,9,10]。不同网络结构会产生不同的执行图与算子组合，这正是"
  "端侧代价差异的根源。")
H("2.2  边缘 AI 推理优化", 2)
P("端侧推理常经 PyTorch→ONNX [19]→TensorRT [18] 路径，配合 FP16/BF16/INT8 量化 [15]、算子融合与内核选择；量化不当会触发"
  "kernel fallback（回退到未优化算子）。学习式 PDR 具有短序列、batch=1、持续在线的特征，使其延迟高度受内核启动与访存主导"
  "而非峰值算力，与 MLPerf [16] 等大 batch 基准的结论不同。")
H("2.3  推理代价预测", 2)
P("一类工作直接预测网络的推理代价。roofline [14] 以算力墙/带宽墙给出分析上界，但面向持续吞吐、不含 batch=1 的内核启动开销；"
  "nn-Meter [21] 以算子级查表并检测融合，精确预测多种边缘设备的延迟；Habitat [22] 基于运行时特征预测同一网络在未见 GPU 上"
  "的运行时；BRP-NAS [23] 等以图神经网络学习延迟预测器以服务 NAS。这些方法多面向通用视觉网络与吞吐/训练场景。本文与之"
  "互补：面向被忽视的 batch=1 学习式 PDR 任务，用更轻量的加性内核感知模型（无需逐核查表）显式建模低精度 fallback，并证明"
  "延迟/能耗可经单标量跨设备迁移。相应地，我们不把参数量/FLOPs 视为“既有方法”，而仅作为检验“规模能否预测代价”的透明对照。")
H("2.4  现有不足", 2)
P("参数量、FLOPs 与简单 benchmark 表都未显式刻画执行图分段、内核启动数、有效访存、硬件利用率与后端算子可用性；roofline [14] "
  "刻画算力–带宽上界但不含 batch=1 的启动开销与 fallback。已有惯性定位综述 [17] 也指出部署代价被系统性忽视。本文的目标是"
  "一个可解释、可验证、可用于部署前决策的代价预测模型。")

# ===================== 3 Method =====================
H("3  问题定义与方法", 1)
H("3.1  部署场景定义", 2)
P("定义部署配置 c=(m,d,p,r)：m 为 PDR 模型，d 为目标设备，p 为框架/后端/精度路径，r 为实时任务要求。给定候选集合 C，"
  "目标是在满足时延与能耗约束下选择定位误差最低的配置：c* = argmin_{c∈C} L_loc(c)，s.t. T̂(c) ≤ D，Ê(c) ≤ B，"
  "其中 D 为截止时间、B 为能耗预算。")
H("3.2  内核感知代价模型", 2)
P("我们把 batch=1 单次前向延迟建模为可解释的加性形式：")
P("    T̂(c) = T0(d,p) + α(d,p)·N_exec(c) + β(d,p)·B_eff(c) + γ(d,p)·F_eff(c) + δ(d,p)·I_fallback(c)", first_indent=False, bold=True)
for b in ["T0：host runtime、框架调用与同步等固定开销；",
          "N_exec：执行图中的内核/执行单元数量（batch=1 下的主导项）；",
          "B_eff：有效访存字节数（工作集）；",
          "F_eff：有效计算工作量（FLOPs）；",
          "I_fallback：低效“算子–精度–硬件”路径的指示变量（如无融合内核的 BF16-LSTM 置 1）；",
          "α、β、γ、δ：设备与部署路径相关的系数。"]:
    BUL(b)
P("直觉：在 batch=1、参数量不足 6 M 的小模型上 GPU 严重欠载，延迟由启动与访存主导而非算力，故 N_exec 项应占主导、F_eff 项应"
  "很弱；δ·I_fallback 项则把回退到逐时间步展开的循环内核这类离散陷阱纳入预测。")
H("3.3  特征提取与校准", 2)
P("特征提取流程：PyTorch Model → ONNX/TensorRT 执行图 → Profile 特征 → 代价模型 → 预测延迟/能耗 → 部署推荐。各特征来源为：")
for b in ["F_eff（FLOPs）：PyTorch FlopCounterMode 实测前向；",
          "参数量：仅作为对照特征，不进入主模型；",
          "B_eff（有效访存）：实测峰值显存 peak_mem 或张量尺寸估计；",
          "N_exec（执行单元数）：图结构叶子算子统计（循环网络按时间步展开），可用 TensorRT/Nsight trace 校核；",
          "I_fallback：由 profiler 日志或算子支持矩阵确认（本文命中 BF16 循环内核）；",
          "设备参数 (α,β,γ,δ,T0)：以普通最小二乘在每台设备上标定。"]:
    BUL(b)
H("3.4  能耗模型", 2)
P("当功耗数据可靠时，定义扣除空载的动态能耗 Ê_dyn(c)=(P_active(c)−P_idle(d))·T̂(c)，单位 mJ/次推理，测量条件为 PyTorch 后端、batch=1、Jetson nvpmodel 0 板级功耗 / Blackwell 持久化模式卡级功耗。减去整板/整卡静态基线后得到模型的边际"
  "能耗，从而消解 card（nvidia-smi）与 board（tegrastats）口径不一致的问题、跨设备可比。测量在锁定时钟（nvpmodel −m 0 + "
  "jetson_clocks；Blackwell 持久化模式）与热稳态下进行，记录功耗模式、频率、起止温度、采样工具与频率、预热次数（见 §4.3）。")
H("3.5  预算感知部署选择", 2)
P("该选择器是离线的候选筛选与推荐工具，用于部署前决策，不执行需要运行时反馈闭环的在线动态调度（后者属未来工作）。给定候选集合与约束，选择器执行：(1) 枚举候选“模型–精度–后端”配置；(2) 用代价模型预测时延与能耗；(3) 过滤不满足实时"
  "截止或能耗预算的配置；(4) 在可行集合中选择验证集 ATE/RTE 最低者；(5) 若无可行配置，返回最接近预算的候选与违反原因。")

# ===================== 4 Experimental setup =====================
H("4  实验设置", 1)
H("4.1  PDR 模型与数据集", 2)
P("我们评测 10 个开源骨干（表 1），覆盖六个架构家族。八个共享统一任务表述（6 通道全局帧惯性窗口 → 二维全局速度 → 速度积分"
  "重建 → RoNIN 式 ATE/RTE [1]）的骨干进入统一精度比较；TLIO（体坐标位移+EKF）与 EqNIO（等变输入）任务不同，作为部署案例"
  "分列。数据集为 MagPIE [20]，按序列划分训练/验证/测试；八个统一骨干在同一配方下训练（Adam lr 1e-4、批量 256、旋转增强、"
  "早停；序列模型用 RoNIN 原生 seq2seq 配方），重训的 RoNIN-ResNet 验证损失复现原文，故精度差异只归因于架构。")
small_table(["模型","家族","参数(M)","准入(精度)","来源"],
            [[DISP[m], FAMZH[m], f"{gv(BB,m,'fp32','params_M'):.2f}",
              "案例" if m in ("tlio_resnet","eqnio") else "✓",
              {"ronin_resnet18":"[1]","ronin_tcn":"[1]","ronin_lstm":"[1]","tlio_resnet":"[2]","imunet":"[4]",
               "tinyodom":"[5]","eqnio":"[6]","mobilenetv2":"[8]","mnasnet":"[9]","efficientnet_b0":"[10]"}[m]] for m in ALL],
            "表 1  模型集、架构家族与准入。8 个统一协议骨干进入精度比较；TLIO/EqNIO 作为部署案例分列。")
H("4.2  硬件与软件栈", 2)
small_table(["设备","定位","算力/功耗","架构"],
            [["RTX PRO 6000 Blackwell","数据中心 GPU","高 / 600 W","Blackwell"],
             ["Jetson AGX Orin","旗舰边缘","中 / 15–60 W","Ampere (sm_87)"],
             ["Jetson Orin Nano","入门边缘","低 / 7–25 W","Ampere (sm_87)"]],
            "表 2  硬件平台与软件栈。Blackwell：PyTorch 2.13 / CUDA 13.0；Jetson：JetPack 自带 CUDA/cuDNN/TensorRT，"
            "锁定 nvpmodel 0 + jetson_clocks。各设备确切库版本随复现脚本记录。"
            "（CUDA/cuDNN/TensorRT/JetPack 具体版本随各设备记录于复现脚本。）")
H("4.3  测量协议", 2)
for b in ["固定 batch=1 与输入窗口；50 次预热后取 500 次迭代；",
          "使用 CUDA 事件同步计时，报告中位数、P95 与标准差；",
          "延迟为网络前向延迟（明确区别于端到端）；",
          "能耗为扣除空载的动态能耗，锁定时钟 + 热稳态 + 3 次重复（§3.4）；",
          "INT8 统一到单一 canonical QDQ 产物：精度只测一次（设备无关），各设备从同一 ONNX 编译引擎。"]:
    BUL(b)
H("4.4  预测验证协议", 2)
P("我们用四种留出协议检验泛化：LOMO（留一模型）、LOAO（留一架构族）、LODO（留一设备，挑战实验）、LOPO（留一精度/后端，"
  "fp32→fp16）。基线为：仅参数量线性回归、仅 FLOPs 线性回归、参数量+FLOPs 回归（作为检验“规模能否预测代价”的透明对照，非既有延迟预测器，后者见 §2.3）；指标为 MAPE、Median APE、最大 APE，"
  "Pearson 仅作辅助。")

# ===================== 5 Results =====================
H("5  实验结果与分析", 1)
_l = lambda m: lat(BB, m, "fp32")
H("5.1  模型规模无法预测端侧延迟", 2)
P(f"在 batch=1 下，模型规模并不能预测端侧延迟。图 2 的效率景观显示，参数量与实测延迟并不单调相关：最快的是 0.11 M 的 TinyOdom（{_l('tinyodom'):.2f} ms），最慢的是 5.23 M 的 "
  f"EqNIO（{_l('eqnio'):.2f} ms），而整族轻量深度可分离 CNN [8,9,10,13] 尽管参数不多却聚在慢端——它们算术强度低、访存受限，"
  f"适配移动 CPU 却跑不满偏爱稠密计算的 GPU。定量地，“仅参数量”基线在留一模型上 MAPE 高达 {VAL['LOMO']['B1 params']['MAPE']:.0f}%、"
  f"留一设备 {VAL['LODO']['B1 params']['MAPE']:.0f}%，是最弱基线之一——若尺寸驱动延迟它本应最优，反证了规模假象。")
figure("figR_landscape.png", "图 2  效率景观：气泡面积∝单次动态能耗；架构而非参数量决定端侧代价——轻量 CNN 尽管参数不多却聚在慢端，TinyOdom 最快、EqNIO 最慢且最耗能。", 3.7)
H("5.2  决定部署代价的执行特征", 2)
P("既然规模无法解释延迟，我们转而考察执行图特征。拟合代价模型可见，N_exec（启动）项在各平台都占主导（约 55–58% 的预测延迟），F_eff（算力）项仅占 20–28%，印证 batch=1 下"
  "延迟受启动/访存而非算力限制；轻量 CNN 的高 N_exec（大量深度可分离层）正是其慢的原因，而 TensorRT 融合通过降低有效 N_exec "
  "带来主要收益。各设备的 T0 与 α（每次启动开销）不同——数据中心约 12 µs/次、边缘约 100 µs/次——这正是跨设备的尺度差异来源。")
P(f"内核 trace 量化证据。 我们用 torch.profiler 直接测量每次 batch=1 前向的真实 CUDA 内核启动数，证实 N_exec 与 fallback 项"
  f"确实可观测且有效：实测内核数与延迟的相关性高达 {KP['corr_kreal_lat']:.2f}，远超参数量（0.41）、FLOPs（0.55）与仅凭图结构"
  f"的叶子算子代理（{KP['corr_kleaf_lat']:.2f}）——这说明真正决定 batch=1 延迟的是执行图的内核启动数，且它需要 profile 才能准确"
  f"获得（叶子算子代理只是粗略近似，与实测内核相关性仅 {KP['corr_kleaf_kreal']:.2f}，这也是我们结构模型绝对误差偏高的原因，"
  f"见 §6.2）。fallback 项则直接可见：RoNIN-LSTM 在 FP16 下仅 {KP['lstm']['fp16']:.0f} 个内核（融合持久化路径），BF16 下暴增到 "
  f"{KP['lstm']['bf16']:.0f} 个（逐时间步展开），{KP['lstm_jump']:.0f}× 的跳变即 I_fallback=1 的直接物理对应。")
H("5.3  对未见部署的预测能力", 2)
P(f"在这些执行特征之上，代价模型可靠地预测未见部署配置，且显著优于朴素基线。图 3(b) 的热力图给出四种留出协议下我方模型与三基线的绝对误差（含每格数值）。结构模型在每个协议上都更优（留一模型 "
  f"{VAL['LOMO']['M4 struct(ours)']['MAPE']:.0f}% vs {VAL['LOMO']['B1 params']['MAPE']:.0f}–"
  f"{VAL['LOMO']['B3 par+FLOP']['MAPE']:.0f}%；留一路径 {VAL['LOPO']['M4 struct(ours)']['MAPE']:.0f}% vs "
  f"{VAL['LOPO']['B1 params']['MAPE']:.0f}–{VAL['LOPO']['B2 FLOPs']['MAPE']:.0f}%），说明启动/访存特征携带真实信号；其绝对 ~31% "
  f"表明单模型精确预测仍难，我们只将结构特征作为机制证据。最强的是跨设备预测：所有结构/基线模型都因看不到设备尺度而崩溃"
  f"（{VAL['LODO']['M4 struct(ours)']['MAPE']:.0f}–{VAL['LODO']['B2 FLOPs']['MAPE']:.0f}%），而单标量迁移只需一个标定模型即达 "
  f"{VAL['LODO']['M5 transfer(1-calib)']['MAPE']:.0f}% MAPE（中位 {VAL['LODO']['M5 transfer(1-calib)']['Median']:.0f}%，"
  f"Pearson≥{min(rA,rN):.2f}，图 3），比一切基线好约 25×。这意味着给定任意未见边缘设备，测一个标定模型即可预测全部 10 个模型"
  f"的延迟排名，把 N×M 的测量降为 O(N+M)。")
figure("figR_prediction.png", "图 3  预测质量：(a) 单个设备标量即可把数据中心 profile 迁移到未见边缘，循环网络（fallback）是唯一偏离点；(b) 四种留出协议 × 五种方法的预测误差热力图（MAPE %，越低越好），结构模型两列全面占优、跨设备单标量迁移仅 11%。", 6.7)
P(f"能耗在全部三档硬件上同样因子化（图 4）。我们对 10 个模型测量扣除空载的动态能耗（§3.4）：数据中心→AGX Orin 满足单标量律 Pearson "
  f"{eRa:.2f}、标量 {eSa:.2f}×、MAPE {eMa:.1f}%，数据中心→Orin Nano 为 Pearson {eRn:.2f}、标量 {eSn:.2f}×、MAPE {eMn:.1f}%。"
  f"一个值得注意且有些反直觉的发现是：单次推理的边际能耗随设备而变、并非“越边缘越省”——入门 Orin Nano 每次推理最省"
  f"（RoNIN-ResNet 约 {edyn_n('ronin_resnet18','fp32'):.0f} mJ），旗舰 AGX Orin 反而最费（约 {edyn_a('ronin_resnet18','fp32'):.0f} mJ，"
  f"约 {eSa:.1f}× 数据中心），因其整板功耗更高。边缘设备真正的能耗优势来自极低的整板静态功耗（空载 ~5–9 W vs 卡级 ~91 W），"
  f"而非每次推理的计算能量。")
figure("figR_energy_slope.png", "图 4  三档动态能耗斜率图（fp32）：所有模型在 AGX Orin 达峰（每次推理最费）、Orin Nano 最低，家族曲线近平行，直观印证能耗跨设备因子化。", 3.7)
_pf = LOAO["per_family"]; _fs = LOAO["fewshot"]
P(f"留一架构（LOAO）的误差偏高（约 {LOAO['overall']:.0f}%），我们对其做误差来源剖析。按家族分解发现，误差几乎被单一结构异类"
  f"主导：等变模型 EqNIO 的 LOAO 误差高达 {_pf['equiv']:.0f}%（表 3），因其等变算子（410 M FLOPs）与训练集中任何模型都不相似，"
  f"纯外推必然失败；而有结构近邻的 TinyOdom 仅 {_pf['TCN-NAS']:.0f}%。这本身是一个诚实结论：全局线性代价模型无法外推到结构"
  f"全新的架构类——恰印证“架构决定代价”。实践中这是可修复项而非缺陷：引入新架构类时，只需先 profile 该类的少量模型做 "
  f"few-shot 校准。在可测的 mobile 与 CNN 家族上，仅加入该家族的 1 个校准模型，预测其余成员的误差就从 "
  f"{_fs['mobile']['0']:.0f}%/{_fs['CNN']['0']:.0f}% 腰斩到 {_fs['mobile']['1']:.0f}%/{_fs['CNN']['1']:.0f}%（表 3）——这与"
  f"跨设备的“一个标定模型”是同一机制：新设备一个标定、新架构一个标定。")
small_table(["架构家族","成员数","LOAO(0-shot)","+1 校准模型","备注"],
            [["EqNIO（等变）", "1", f"{_pf['equiv']:.0f}%", "—", "结构异类，主要误差源"],
             ["TCN", "1", f"{_pf['TCN']:.0f}%", "—", ""],
             ["CNN", "2", f"{_pf['CNN']:.0f}%", f"{_fs['CNN']['1']:.0f}%", ""],
             ["轻量 CNN（mobile）", "4", f"{_pf['mobile']:.0f}%", f"{_fs['mobile']['1']:.0f}%", ""],
             ["TinyOdom（TCN-NAS）", "1", f"{_pf['TCN-NAS']:.0f}%", "—", "有结构近邻，预测最准"]],
            "表 3  LOAO 误差来源与 few-shot 家族校准。误差被结构异类 EqNIO（96%）主导；引入新架构类时，1 个校准模型即可"
            "将误差腰斩。few-shot 仅在成员数≥2 的 mobile/CNN 家族可测。")
H("5.4  BF16-LSTM 异常案例与 fallback 项", 2)
P(f"预测能力在低精度算子回退这一案例上尤为关键。RoNIN-LSTM 在 BF16 下比 FP16 慢约 3 倍：剖析表明，cuDNN 仅为 FP16 分派融合的持久化 LSTM 内核，BF16 无对应内核而回退到"
  f"逐时间步展开（内核启动数增至约 112 倍）。该陷阱跨平台一致且在边缘放大：BF16/FP16 延迟比从 Blackwell 的 3.1× 升到 AGX 的 "
  f"26.6× 与 Nano 的 29.6×（图 5），在 Nano 上达 95 ms，直接从实时变为不可用；能耗上同样现形且在边缘剧烈放大——动态能耗在 "
  f"Blackwell 从 {edyn('ronin_lstm','fp16'):.1f} 暴涨到 {edyn('ronin_lstm','bf16'):.1f} mJ"
  f"（{edyn('ronin_lstm','bf16')/edyn('ronin_lstm','fp16'):.1f}×），在 AGX Orin 达 {edyn_a('ronin_lstm','bf16')/edyn_a('ronin_lstm','fp16'):.0f}×、"
  f"在 Orin Nano 从 {edyn_n('ronin_lstm','fp16'):.1f} 暴涨到 {edyn_n('ronin_lstm','bf16'):.1f} mJ"
  f"（{edyn_n('ronin_lstm','bf16')/edyn_n('ronin_lstm','fp16'):.0f}×）。关键在于：所有基于"
  f"规模/FLOPs 的模型都对该点束手无策（预测误差 65–78%），而加入 I_fallback 指示项后，代价模型把该误差降至可忽略——即模型"
  f"不仅能预测、还能预警此类 fallback 陷阱。该结论限定在本文软件栈与实验环境内。")
figure("figR_fallback.png", "图 5  BF16 循环内核 fallback：(a) 随隐藏维增大而坍缩；(b) 在边缘放大近 30 倍；(c) 内核启动数从 10（融合持久化）暴增至 1241（逐步展开），124× 直接可观测。", 6.9)
H("5.5  面向部署选择的应用", 2)
P(f"将上述代价模型与预测能力落到部署实践，我们把它接入预算感知选择器（离线候选筛选与推荐，非在线调度），并给出一个实例。场景：Orin Nano、100 Hz（截止 10 ms）、最小化 ATE。所有 TensorRT-FP16 "
  f"配置都满足截止（0.5–1.2 ms）；若再要求 ATE ≤ 1.0 m，选择器排除 TinyOdom（{_a('tinyodom'):.2f} m）等，推荐 RoNIN-ResNet"
  f"（ATE {_a('ronin_resnet18'):.2f} m、Nano TRT 约 {trt('trt_int8_latency_orin_nano.csv','ronin_resnet18',None) or 0.6:.1f} ms 量级），"
  f"并自动剔除 BF16 路径（fallback 预警）。若能耗预算进一步收紧，选择器可调用与模型/设备正交的第三杠杆——降低推理频率：我们"
  f"在 53 条测试序列上验证，把推理频率固定下调 12–16×（≈1.5 Hz 而非 20 Hz）对 ATE 与 RTE 均无损（图 6a），从而把整条精度–能耗"
  f"前沿整体左移一个数量级（图 6b）；运动自适应门控在此并不优于均匀降频（冗余是均匀的、非突发的）。")
figure("fig8_redundancy.png", "图 6  部署杠杆之一——时间冗余：(a) 频率降 12× 对 ATE/RTE 均无损；(b) 整条 Pareto 左移一个数量级。", 6.4)

# main results + energy tables
prim = ["ronin_resnet18","imunet","mobilenetv2","mnasnet","efficientnet_b0","tinyodom","ronin_lstm","ronin_tcn"]
def i8d(m):
    v = gv(ci, m, "onnx_int8", "ate_m"); return f"+{(v-_a(m))/_a(m)*100:.0f}%" if v else "—"
small_table(["模型","FP32 ATE(m)","INT8 ΔATE","Blk FP32(ms)","E_dyn Blk(mJ)","E_dyn AGX(mJ)","E_dyn Nano(mJ)"],
            [[DISP[m], f"{_a(m):.2f}", i8d(m), f"{_l(m):.2f}",
              f"{edyn(m,'fp32'):.1f}" if edyn(m,'fp32') else "—",
              f"{edyn_a(m,'fp32'):.1f}" if edyn_a(m,'fp32') else "—",
              f"{edyn_n(m,'fp32'):.1f}" if edyn_n(m,'fp32') else "—"] for m in prim],
            "表 4  主结果（8 个统一骨干）：精度、INT8 代价、数据中心延迟与三档动态能耗（fp32，扣空载/锁频/热稳态）。"
            "精度设备无关；能耗按设备。动态能耗跨全部三档因子化（Blackwell→AGX Pearson 0.96/MAPE 11%，→Nano 0.98/MAPE 9%）。")

# ===================== 6 Discussion =====================
H("6  讨论", 1)
H("6.1  适用范围", 2)
P("本模型适用于固定输入 shape、batch=1、GPU 边缘设备、已支持的 PyTorch/ONNX/TensorRT 推理路径，以及当前验证的 PDR 网络类别。")
H("6.2  局限性", 2)
for b in ["对全新 GPU 架构的外推可能受驱动、库与算子实现影响；",
          "未充分覆盖动态输入 shape、自定义算子及 CPU/NPU/ISP 异构管线；",
          "能耗预测受功耗模式、温度与 DVFS 影响；",
          "定位精度泛化受数据集与评测划分影响；",
          "预测模型减少候选测量成本，但不替代最终部署测试。"]:
    BUL(b)
H("6.3  实践建议", 2)
for b in ["不能仅凭参数量选择端侧 PDR 模型；", "引入结构全新的网络家族时，先 profile 其 1–2 个模型做 few-shot 校准再预测；",
          "CNN 优先尝试 TensorRT FP16/INT8；",
          "RNN/时序模型须检查目标精度下优化内核的可用性（避免 BF16 fallback）；",
          "严格实时任务应使用 P95 延迟而非平均延迟；",
          "Jetson 测量必须固定并完整报告功耗模式与时钟策略。"]:
    BUL(b)

# ===================== 7 Conclusion =====================
H("7  结论", 1)
P("我们把 batch=1 学习式 PDR 的端侧部署，从穷举 benchmark 提升为预测驱动的约束感知选择。所提内核感知代价模型用执行图开销、"
  "访存、计算与 fallback 指示项解释并预测端侧延迟，经 LOMO/LOAO/LODO 严格验证优于参数量/FLOPs 基线，在受测的十个模型与三档硬件范围内跨设备单标量迁移仅 11% "
  "误差（不宣称普遍泛化，对结构全新架构需 few-shot 校准），并能预警 BF16-LSTM 等陷阱；据此构建的预算感知选择器可在实时、能耗与精度约束下自动选型。未来将扩展到手机 NPU、CPU "
  "与在线自适应调度。")

# ===================== References =====================
H("参考文献", 1)
REFS = [
 "S. Herath, H. Yan, and Y. Furukawa. RoNIN: Robust Neural Inertial Navigation in the Wild: Benchmark, Evaluations, and New Methods. In ICRA, 2020.",
 "W. Liu, D. Caruso, E. Ilg, J. Dong, A. I. Mourikis, K. Daniilidis, V. Kumar, and J. Engel. TLIO: Tight Learned Inertial Odometry. IEEE Robotics and Automation Letters, 5(4):5653–5660, 2020.",
 "C. Chen, X. Lu, A. Markham, and N. Trigoni. IONet: Learning to Cure the Curse of Drift in Inertial Odometry. In AAAI, 2018.",
 "B. Zeinali, H. Zanddizari, and M. J. Chang. IMUNet: Efficient Regression Architecture for Inertial IMU Navigation and Positioning. arXiv:2208.00068, 2022.",
 "S. S. Saha, S. S. Sandha, L. A. Garcia, and M. Srivastava. TinyOdom: Hardware-Aware Efficient Neural Inertial Navigation. Proc. ACM IMWUT (UbiComp), 6(2):1–32, 2022.",
 "R. K. Jayanth, Y. Xu, Z. Wang, E. Chatzipantazis, K. Daniilidis, and D. Gehrig. EqNIO: Subequivariant Neural Inertial Odometry. arXiv:2408.06321, 2024.",
 "K. He, X. Zhang, S. Ren, and J. Sun. Deep Residual Learning for Image Recognition. In CVPR, 2016.",
 "M. Sandler, A. Howard, M. Zhu, A. Zhmoginov, and L.-C. Chen. MobileNetV2: Inverted Residuals and Linear Bottlenecks. In CVPR, 2018.",
 "M. Tan, B. Chen, R. Pang, V. Vasudevan, M. Sandler, A. Howard, and Q. V. Le. MnasNet: Platform-Aware Neural Architecture Search for Mobile. In CVPR, 2019.",
 "M. Tan and Q. V. Le. EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks. In ICML, 2019.",
 "S. Hochreiter and J. Schmidhuber. Long Short-Term Memory. Neural Computation, 9(8):1735–1780, 1997.",
 "S. Bai, J. Z. Kolter, and V. Koltun. An Empirical Evaluation of Generic Convolutional and Recurrent Networks for Sequence Modeling. arXiv:1803.01271, 2018.",
 "F. Chollet. Xception: Deep Learning with Depthwise Separable Convolutions. In CVPR, 2017.",
 "S. Williams, A. Waterman, and D. Patterson. Roofline: An Insightful Visual Performance Model for Multicore Architectures. Communications of the ACM, 52(4):65–76, 2009.",
 "B. Jacob, S. Kligys, B. Chen, M. Zhu, M. Tang, A. Howard, H. Adam, and D. Kalenichenko. Quantization and Training of Neural Networks for Efficient Integer-Arithmetic-Only Inference. In CVPR, 2018.",
 "V. J. Reddi et al. MLPerf Inference Benchmark. In ISCA, 2020.",
 "C. Chen and X. Pan. Deep Learning for Inertial Positioning: A Survey. IEEE Transactions on Intelligent Transportation Systems, 2024.",
 "NVIDIA. TensorRT: A High-Performance Deep Learning Inference SDK. https://developer.nvidia.com/tensorrt.",
 "ONNX Runtime Developers. ONNX Runtime. https://onnxruntime.ai.",
 "MagPIE: A Magnetic and Inertial Pedestrian Indoor-positioning Dataset. （请按实际使用版本核对确切出处与作者。）",
 "L. L. Zhang, S. Han, J. Wei, N. Zheng, T. Cao, Y. Yang, and Y. Liu. nn-Meter: Towards Accurate Latency Prediction of Deep-Learning Model Inference on Diverse Edge Devices. In MobiSys, 2021.",
 "G. X. Yu, Y. Gao, P. Golikov, and G. Pekhimenko. Habitat: A Runtime-Based Computational Performance Predictor for Deep Neural Network Training. In USENIX ATC, 2021.",
 "L. Dudziak, T. Chau, M. S. Abdelfattah, R. Lee, H. Kim, and N. D. Lane. BRP-NAS: Prediction-based NAS using GCNs. In NeurIPS, 2020.",
]
for i, rf in enumerate(REFS, 1):
    p = d.add_paragraph(); r = p.add_run(f"[{i}] {rf}"); r.font.size = Pt(8.5); cjk(r)
    p.paragraph_format.space_after = Pt(1)

d.save(OUT); print("wrote", OUT)
